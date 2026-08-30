from __future__ import annotations

import json
import shutil
import subprocess
from collections import Counter
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from aisurgeon_decentralised.drug_extraction_schema import (
    DrugCatalogEntry,
    DrugMention,
    SourceRecord,
)


def read_jsonl(path: Path, model):
    rows = []
    with path.open("r", encoding="utf-8") as handle:
        for line in handle:
            if line.strip():
                rows.append(model.model_validate_json(line))
    return rows


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text or "")
    run.bold = bold
    run.font.size = Pt(8.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def add_compact_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header_row = table.rows[0]
    set_repeat_table_header(header_row)
    for idx, header in enumerate(headers):
        set_cell_text(header_row.cells[idx], header, bold=True)
        set_cell_shading(header_row.cells[idx], "D9EAF7")
        header_row.cells[idx].width = Cm(widths[idx])
    for row_values in rows:
        row = table.add_row()
        for idx, value in enumerate(row_values):
            set_cell_text(row.cells[idx], value)
            row.cells[idx].width = Cm(widths[idx])
    document.add_paragraph()


def add_heading(document: Document, text: str, level: int = 1) -> None:
    document.add_heading(text, level=level)


def joined(values: list[str], max_items: int = 8) -> str:
    if not values:
        return "-"
    shown = values[:max_items]
    suffix = "" if len(values) <= max_items else f"; +{len(values) - max_items} weitere"
    return "; ".join(shown) + suffix


def pages(entry: DrugCatalogEntry) -> str:
    parts = []
    for source, page_list in sorted(entry.exact_pdf_pages_by_source.items()):
        collapsed = ", ".join(str(page) for page in page_list[:18])
        if len(page_list) > 18:
            collapsed += f", +{len(page_list) - 18}"
        parts.append(f"{source}: {collapsed}")
    return "\n".join(parts)


def evidence_summary(entry: DrugCatalogEntry) -> str:
    if not entry.evidence:
        return "-"
    evidence = entry.evidence[0]
    return f"{evidence.source_id} S. {evidence.pdf_page}: \"{evidence.exact_context_quote}\""


def clinical_summary(entry: DrugCatalogEntry) -> str:
    contexts = [context for context in entry.clinical_contexts if context]
    return joined(contexts, max_items=3)


def entry_rows(entries: list[DrugCatalogEntry]) -> list[list[str]]:
    return [
        [
            entry.canonical_name_de,
            entry.canonical_name_en or "-",
            joined(entry.raw_mentions, 6),
            joined(entry.drug_classes, 4),
            joined(entry.regimen_names, 4),
            ", ".join(entry.guideline_sources),
            pages(entry),
            clinical_summary(entry),
            evidence_summary(entry),
            entry.citation_verification_status,
        ]
        for entry in entries
    ]


def source_rows(sources: list[SourceRecord]) -> list[list[str]]:
    return [
        [
            source.source_id,
            source.source_filename,
            source.document_status,
            str(source.pdf_pages),
            source.source_sha256,
        ]
        for source in sources
    ]


def add_summary(document: Document, report: dict[str, Any], mentions: list[DrugMention]) -> None:
    add_heading(document, "Zusammenfassung", 1)
    mention_counts = Counter(mention.source_id for mention in mentions)
    rows = [
        ["Eindeutige konkrete Wirkstoffe", str(report["unique_active_substances"])],
        ["Regime", str(report["regimens"])],
        ["Wirkstoffgruppen", str(report["drug_classes"])],
        ["Handelsnamen", str(report["brand_names"])],
        ["PRIORITY_A", str(report["priority_counts"]["PRIORITY_A"])],
        ["PRIORITY_B", str(report["priority_counts"]["PRIORITY_B"])],
        ["PRIORITY_C", str(report["priority_counts"]["PRIORITY_C"])],
        ["MANUAL_REVIEW", str(report["priority_counts"]["MANUAL_REVIEW"])],
    ]
    for source_id, count in sorted(mention_counts.items()):
        rows.append([f"Einzelnennungen {source_id}", str(count)])
    add_compact_table(document, ["Kennzahl", "Wert"], rows, [8.0, 7.0])


def set_document_style(document: Document) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.3)
    section.bottom_margin = Cm(1.3)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "Arial"
        styles[style_name].font.bold = True


def build_docx(output_dir: Path) -> Path:
    sources = read_jsonl(output_dir / "source_registry.jsonl", SourceRecord)
    mentions = read_jsonl(output_dir / "drug_mentions.jsonl", DrugMention)
    catalog = read_jsonl(output_dir / "drug_catalog.jsonl", DrugCatalogEntry)
    report = json.loads((output_dir / "extraction_report.json").read_text(encoding="utf-8"))

    document = Document()
    set_document_style(document)

    document.add_heading("Medikamente und Wirkstoffe in den AISurgeon-Pilotleitlinien", 0)
    p = document.add_paragraph()
    p.add_run(f"Erstellungsdatum: {report['created_at']}\n")
    p.add_run(f"extraction_run_id: {report['extraction_run_id']}")

    add_heading(document, "Methodischer Hinweis", 1)
    document.add_paragraph(
        "Die native PDF-Analyse erfolgte mit Gemini 3.5 Flash über alle drei Leitlinien "
        "in deterministischen Seitenbatches mit maximal 20 Original-PDF-Seiten und einer Seite Überlappung. "
        "Klinische Inhalte wurden von rein bibliografischen Erwähnungen getrennt. "
        "Die Normalisierung erfolgte ohne externe Arzneimitteldatenbank. "
        "Eine Nennung stellt keine Empfehlung dar."
    )

    add_heading(document, "Quellenregister", 1)
    add_compact_table(
        document,
        ["source_id", "Dateiname", "Dokumentstatus", "PDF-Seiten", "SHA-256"],
        source_rows(sources),
        [4.0, 9.0, 4.0, 2.0, 8.0],
    )

    add_summary(document, report, mentions)

    priority_a = [entry for entry in catalog if entry.priority == "PRIORITY_A"]
    priority_b = [entry for entry in catalog if entry.priority == "PRIORITY_B"]
    priority_c = [entry for entry in catalog if entry.priority == "PRIORITY_C"]
    manual = [entry for entry in catalog if entry.priority == "MANUAL_REVIEW"]
    regimens = [entry for entry in catalog if entry.entity_type == "COMBINATION_REGIMEN"]
    classes = [entry for entry in catalog if entry.entity_type == "DRUG_CLASS"]

    headers = [
        "Kanonischer Name",
        "EN/INN",
        "Rohformen",
        "Klasse",
        "Regime",
        "Leitlinien",
        "PDF-Seiten",
        "Kontext",
        "Belegstelle",
        "Status",
    ]
    widths = [3.0, 2.3, 3.0, 2.6, 2.4, 2.3, 2.8, 4.0, 5.0, 2.0]

    add_heading(document, "PRIORITY_A", 1)
    add_compact_table(document, headers, entry_rows(priority_a), widths)

    add_heading(document, "PRIORITY_B", 1)
    add_compact_table(document, headers, entry_rows(priority_b), widths)

    add_heading(document, "Kombinationstherapien und Regime", 1)
    document.add_paragraph(
        "Regime sind getrennt von einzelnen Wirkstoffen dargestellt. Komponenten gelten nur dann als source-explicit, "
        "wenn sie in den PDF-Belegen ausgeschrieben oder eindeutig zugeordnet wurden."
    )
    add_compact_table(document, headers, entry_rows(regimens), widths)

    add_heading(document, "Nur Als Wirkstoffgruppe Genannte Begriffe", 1)
    add_compact_table(document, headers, entry_rows(classes), widths)

    add_heading(document, "PRIORITY_C / Ausschließlich Literaturverzeichnis", 1)
    add_compact_table(document, headers, entry_rows(priority_c), widths)

    add_heading(document, "Manuell Zu Prüfende Einträge", 1)
    manual_rows = []
    for entry in manual:
        manual_rows.append(
            [
                entry.canonical_name_de,
                entry.entity_type,
                joined(entry.raw_mentions, 6),
                pages(entry),
                "; ".join(entry.manual_review_reasons) or "Kontext oder Normalisierung unsicher.",
                evidence_summary(entry),
            ]
        )
    add_compact_table(
        document,
        ["Name", "Typ", "Rohformen", "PDF-Seiten", "Grund", "Belegstelle"],
        manual_rows,
        [4.0, 3.0, 4.5, 4.0, 6.0, 7.0],
    )

    add_heading(document, "Methodische Limitationen", 1)
    for text in [
        "Die HCC/BCC-Leitlinie ist eine Konsultationsfassung.",
        "Eine Nennung bedeutet nicht automatisch eine Empfehlung.",
        "Eine bibliografische Nennung bedeutet nicht aktuelle klinische Relevanz.",
        "In diesem Schritt wurde keine externe Arzneimittel-Normalisierung verwendet.",
        "Lokale Zitatprüfung kann bei PDF-Strukturproblemen Einträge als UNVERIFIED markieren, ohne sie zu löschen.",
    ]:
        document.add_paragraph(text, style=None)

    path = output_dir / "Medikamente_und_Wirkstoffe_aus_drei_Leitlinien.docx"
    document.save(path)
    return path


def render_docx_for_qa(docx_path: Path, qa_dir: Path) -> dict[str, Any]:
    qa_dir.mkdir(parents=True, exist_ok=True)
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return {
            "status": "NOT_AVAILABLE",
            "reason": "LibreOffice/soffice not found; DOCX visual render QA could not be performed.",
            "png_pages": [],
        }

    pdf_dir = qa_dir / "pdf"
    png_dir = qa_dir / "png"
    pdf_dir.mkdir(parents=True, exist_ok=True)
    png_dir.mkdir(parents=True, exist_ok=True)
    subprocess.run(
        [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(pdf_dir), str(docx_path)],
        check=True,
        capture_output=True,
        text=True,
    )
    pdf_path = pdf_dir / (docx_path.stem + ".pdf")
    subprocess.run(
        ["pdftoppm", "-png", "-r", "130", str(pdf_path), str(png_dir / "page")],
        check=True,
        capture_output=True,
        text=True,
    )
    pages = sorted(str(path) for path in png_dir.glob("page-*.png"))
    return {
        "status": "PASS" if pages else "FAIL",
        "reason": None if pages else "No PNG pages rendered.",
        "pdf_path": str(pdf_path),
        "png_pages": pages,
        "inspection_note": "Automated render completed; pages produced for visual inspection.",
    }


def update_report_docx_status(output_dir: Path, qa_result: dict[str, Any]) -> None:
    report_path = output_dir / "extraction_report.json"
    report = json.loads(report_path.read_text(encoding="utf-8"))
    report["docx_render_qa"] = qa_result["status"]
    report["docx_render_qa_detail"] = qa_result
    report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
